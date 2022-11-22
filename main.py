import argparse
from random import randint, shuffle
from re import compile

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

adj_pair_list = [
    {'type': '評価', 'adj_pair': ['立派な', 'ひどい']},
    {'type': '評価', 'adj_pair': ['役立つ', '役立たない']},
    {'type': '評価', 'adj_pair': ['よい', 'わるい']},
    {'type': '評価', 'adj_pair': ['涼しい', '暑い']},
    {'type': '評価', 'adj_pair': ['快適な', '不快な']},
    {'type': '力量', 'adj_pair': ['大きい', '小さい']},
    {'type': '力量', 'adj_pair': ['強い', '弱い']},
    {'type': '力量', 'adj_pair': ['軽い', '重い']},
    {'type': '力量', 'adj_pair': ['柔らかい', '硬い']},
    {'type': '活動', 'adj_pair': ['若い', '老いた']},
    {'type': '活動', 'adj_pair': ['優しい', '厳しい']},
    {'type': '活動', 'adj_pair': ['かっこいい', 'かわいい']},
]


def get_args():
    parser = argparse.ArgumentParser(description='Generate SD-method-sheet from the base Word file.')
    parser.add_argument('-p', '--path', required=True,
                        help='Path to the base Word file')
    parser.add_argument('-s', '--size', type=int, default=1,
                        help='Size of sheets to be generated')
    return parser.parse_args()


if __name__ == '__main__':
    args = get_args()
    # Wordファイルを読み込む
    doc = docx.Document(args.path)
    for i in range(args.size):
        for tbl in doc.tables:
            # 形容詞対の左右をランダムで入れ替える
            for row in adj_pair_list:
                shuffle(row['adj_pair'])
            # 同じtypeの形容詞対が続かないように調整する
            for j in reversed(range(1, len(adj_pair_list))):
                c = 0
                while c < 100:
                    k = randint(0, j)
                    adj_pair_list[j], adj_pair_list[k] = \
                        adj_pair_list[k], adj_pair_list[j]
                    if j < len(adj_pair_list) - 1:
                        if adj_pair_list[j]['type'] == adj_pair_list[j + 1]['type']:
                            c += 1
                            continue
                    break
            # 「よい - わるい」が前半にあったら、後半の要素と入れ替える
            for j in range(len(adj_pair_list)):
                if adj_pair_list[j]['adj_pair'] == ['よい', 'わるい'] \
                        or adj_pair_list[j]['adj_pair'] == ['わるい', 'よい']:
                    if j < len(adj_pair_list) / 2:
                        c = 0
                        while c < 100:
                            k = randint(int(len(adj_pair_list) / 4), len(adj_pair_list) - 1)
                            adj_pair_list[j]['adj_pair'], adj_pair_list[k]['adj_pair'] = \
                                adj_pair_list[k]['adj_pair'], adj_pair_list[j]['adj_pair']
                            if j < len(adj_pair_list) - 1:
                                if adj_pair_list[j]['type'] == adj_pair_list[j + 1]['type']:
                                    c += 1
                                    continue
                            break
                        break

            for j in range(len(adj_pair_list)):
                # 属性・形容詞対を書き込む
                tbl.rows[j + 1].cells[0].text = adj_pair_list[j]['type']
                tbl.rows[j + 1].cells[1].text = adj_pair_list[j]['adj_pair'][0]
                tbl.rows[j + 1].cells[3].text = adj_pair_list[j]['adj_pair'][1]
                # 中央揃えにする
                tbl.rows[j + 1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                tbl.rows[j + 1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                tbl.rows[j + 1].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 属性部分を太字にする
                for run in tbl.rows[j + 1].cells[0].paragraphs[0].runs:
                    if run.text:
                        run.bold = True
            # 列の幅を自動調節する
            tbl.autofit = True

        # Wordファイルを保存する
        pattern = compile(r'.docx')
        path = pattern.sub(f'_{i}.docx', args.path)
        doc.save(path)
        print(f'Create {path}')
    print('Complete!!')
